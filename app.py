import re
import difflib
from io import BytesIO
import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── 스타일 교정 규칙 (패턴, 교체어, 설명) ────────────────────────────────────
# 백레퍼런스는 Python 방식 \1 사용 (JavaScript $1 아님)
RULES = [
    (r'\bunprecedented demands\b',          'stringent requirements',           'unprecedented → stringent'),
    (r'\bremarkable capacity\b',             'outstanding specific capacity',     'remarkable → outstanding'),
    (r'\bplaced unprecedented demands on\b', 'imposed stringent requirements on', 'GPT 미사여구 제거'),
    (r'\bpaving the way for\b',              'providing a pathway for',           '학술적 표현 교정'),
    (r'\bundermines its practical deployment\b', 'limits its practical viability', '학술적 표현 교정'),
    (r'\bculminating in\b',                  'resulting in',                      'culminating → resulting'),
    (r'\bunderscore the need\b',             'highlight the necessity',           '학술적 동사 교체'),
    (r'\bdisproportionate set of\b',         'significant',                       '수식어 압축'),
    (r'\bfading\b',                          'degradation',                       'fading → degradation'),
    (r'\bdeteriorat\w*\b',                   'degraded',                          'deteriorate → degrade'),
    (r'\bdecay\w*\b',                        'degraded',                          'decay → degrade'),
    (r'\bfades?\b',                          'degrades',                          'fade → degrade'),
    (r'\bfaded\b',                           'degraded',                          'faded → degraded'),
    (r'\bfading\b',                          'degrading',                         'fading → degrading'),
    (r'\bdemonstrates\b',                    'exhibits',                          'demonstrates → exhibits'),
    (r'\bdemonstrated\b',                    'exhibited',                         'demonstrated → exhibited'),
    (r'\bdemonstrate\b',                     'exhibit',                           'demonstrate → exhibit'),
    (r'\bshowed\b',                          'exhibited',                         'showed → exhibited'),
    (r'\bshows\b',                           'exhibits',                          'shows → exhibits'),
    (r'\bshow\b',                            'exhibit',                           'show → exhibit'),
    (r'\b(?:is|are) due to\b',               'is attributed to',                  'due to → attributed to'),
    (r'\b(?:is|are) caused by\b',            'is attributed to',                  'caused by → attributed to'),
    (r'\bowing to\b',                        'attributed to',                     'owing to → attributed to'),
    (r'\bthanks to\b',                       'owing to',                          'thanks to → owing to'),
    (r'\bimproved\b',                        'enhanced',                          'improved → enhanced'),
    (r'\bimproves\b',                        'enhances',                          'improves → enhances'),
    (r'\bimprove\b',                         'enhance',                           'improve → enhance'),
    (r'\bboosted\b',                         'enhanced',                          'boosted → enhanced'),
    (r'\bboosts\b',                          'enhances',                          'boosts → enhances'),
    (r'\bboost\b',                           'enhance',                           'boost → enhance'),
    (r'\bAlso,\b',                           'Furthermore,',                      'Also → Furthermore'),
    (r'\bBut,\b',                            'However,',                          'But → However'),
    (r'\bYet\b',                             'However,',                          'Yet → However'),
    (r'\bIn this paper\b',                   'In this work',                      'In this paper → In this work'),
    (r'\bfig\.\b',                           'Figure',                            'fig. → Figure'),
    (r'\bnegative electrode\b',              'anode',                             'negative electrode → anode'),
    (r'\bpositive electrode\b',              'cathode',                           'positive electrode → cathode'),
    (r'\bIn conclusion,',                    'In summary,',                       'In conclusion → In summary'),
]


def _apply_rules(text: str) -> str:
    for pattern, repl, _ in RULES:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
    text = re.sub(r'\b(\w+)\s+based(?!\s+on)\b', r'\1-based', text, flags=re.IGNORECASE)
    return text


def _diff_html(original: str, corrected: str) -> str:
    """단어 단위 diff → 빨간 취소선(원본) + 초록(교정) HTML."""
    orig_w = original.split()
    corr_w = corrected.split()
    matcher = difflib.SequenceMatcher(None, orig_w, corr_w, autojunk=False)
    parts = []
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            parts.append(' '.join(orig_w[i1:i2]))
        elif op == 'replace':
            parts.append(
                f'<del style="color:#c00;text-decoration:line-through;">{" ".join(orig_w[i1:i2])}</del>'
                f' <ins style="color:#060;font-weight:bold;text-decoration:none;">{" ".join(corr_w[j1:j2])}</ins>'
            )
        elif op == 'delete':
            parts.append(f'<del style="color:#c00;text-decoration:line-through;">{" ".join(orig_w[i1:i2])}</del>')
        elif op == 'insert':
            parts.append(f'<ins style="color:#060;font-weight:bold;text-decoration:none;">{" ".join(corr_w[j1:j2])}</ins>')
    return ' '.join(parts)


def process(text: str):
    """Returns (original_en, corrected, diff_html)."""
    if any('가' <= c <= '힣' for c in text[:200]):
        raw = GoogleTranslator(source='ko', target='en').translate(text)
    else:
        raw = text
    corrected = _apply_rules(raw)
    return raw, corrected, _diff_html(raw, corrected)


# ── 첨자 패턴 (6그룹) ────────────────────────────────────────────────────────
# G1 인용번호  G2 단위지수(g-1)  G3 Li+
# G4 부호/글자뒤(Na+, Cl-, COO-)   ← en dash 포함
# G5 이온전하/숫자뒤(SO3-, SO3–)   ← en dash 포함
# G6 화학식 아랫첨자(TiO2)
_SUP_SUB_PAT = re.compile(
    r'(?<=[.!?])(\d+(?:[,\s\-–−]\s*\d+)*)'          # G1
    r'|(?<=[A-Za-z])([-–−]\d+)'                      # G2
    r'|(Li\+)'                                        # G3
    r'|(?<=[A-Za-z])([+\-−–])(?=[^A-Za-z0-9]|$)'   # G4  ← – 추가
    r'|(?<=\d)([+\-−–])(?![A-Za-z0-9])'             # G5  ← – 추가
    r'|(?<=[A-Za-z])(\d+)'                            # G6
)
_PREFIX_DASH = re.compile(r'(?<![A-Za-z0-9])-(?=[A-Z][A-Za-z0-9])')


def _run(paragraph, text, sup=False, sub=False):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if sup: run.font.superscript = True
    if sub: run.font.subscript = True
    return run


def build_docx(text: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    for block in text.split('\n\n'):
        if not block.strip(): continue
        block = _PREFIX_DASH.sub('–', block)
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        pos = 0
        for m in _SUP_SUB_PAT.finditer(block):
            if m.start() > pos: _run(p, block[pos:m.start()])
            g1,g2,g3,g4,g5,g6 = m.groups()
            d = lambda s: s.replace('-','–').replace('−','–')
            if g1:   _run(p, g1, sup=True)
            elif g2: _run(p, d(g2), sup=True)
            elif g3: _run(p, 'Li'); _run(p, '+', sup=True)
            elif g4: _run(p, d(g4), sup=True)
            elif g5: _run(p, d(g5), sup=True)
            elif g6: _run(p, g6, sub=True)
            pos = m.end()
        if pos < len(block): _run(p, block[pos:])
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def web_fmt(text: str) -> str:
    """화학식 첨자를 <sup>/<sub> HTML 태그로 변환."""
    text = _PREFIX_DASH.sub('–', text)
    def _repl(m):
        g1,g2,g3,g4,g5,g6 = m.groups()
        d = lambda s: s.replace('-','–').replace('−','–')
        if g1: return f'<sup>{g1}</sup>'
        if g2: return f'<sup>{d(g2)}</sup>'
        if g3: return 'Li<sup>+</sup>'
        if g4: return f'<sup>{d(g4)}</sup>'
        if g5: return f'<sup>{d(g5)}</sup>'
        if g6: return f'<sub>{g6}</sub>'
        return m.group()
    return _SUP_SUB_PAT.sub(_repl, text)


# ── Streamlit UI ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="논문 스타일 교정기", page_icon="🔋", layout="wide")
st.title("🔋 배터리 논문 스타일 교정 및 번역")

for k in ('orig', 'corr', 'diff'):
    if k not in st.session_state: st.session_state[k] = ''

_CARD = 'font-family:serif;font-size:1.1rem;text-align:justify;line-height:2.0;padding:25px;border-radius:10px;color:#111;'

col_left, col_right = st.columns(2)
with col_left:
    user_input = st.text_area("텍스트 입력 (한글/영어)", height=500)
    if st.button("교정 실행", type="primary", use_container_width=True):
        if user_input.strip():
            with st.spinner("교정 중..."):
                try:
                    orig, corr, diff = process(user_input)
                    st.session_state.orig = orig
                    st.session_state.corr = corr
                    st.session_state.diff = diff
                except Exception as e:
                    st.error(f"오류: {e}")

with col_right:
    if st.session_state.corr:
        tab1, tab2 = st.tabs(["🔴 교정본 (취소선)", "✅ 교정 완료본"])

        with tab1:
            # diff에도 첨자 적용 (교정 완료된 쪽 단어만 web_fmt 통과시킴)
            diff_formatted = web_fmt(st.session_state.diff).replace('\n', '<br>')
            st.markdown(
                f'<div style="{_CARD}background:#fff8f8;border:1px solid #ffcccc;">'
                f'{diff_formatted}</div>',
                unsafe_allow_html=True,
            )

        with tab2:
            st.download_button(
                "📥 Word 다운로드",
                data=build_docx(st.session_state.corr),
                file_name="refined_paper.docx",
                use_container_width=True,
            )
            st.markdown(
                f'<div style="{_CARD}background:#fff;border:1px solid #ddd;">'
                f'{web_fmt(st.session_state.corr).replace(chr(10), "<br>")}</div>',
                unsafe_allow_html=True,
            )
    else:
        st.info("결과가 여기에 표시됩니다.")
