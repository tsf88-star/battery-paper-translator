# -*- coding: utf-8 -*-
import re
import sys
import collections
import os

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CORPUS_PATH = r"C:\Users\cho_b\Documents\이&최 습식 실리콘 논문\merged_corpus.txt"
OUTPUT_PATH = r"C:\Users\cho_b\Documents\이&최 습식 실리콘 논문\논문_스타일_분석보고서.docx"

# ── helpers ───────────────────────────────────────────────────────────────────

def split_sentences(text):
    text = re.sub(r"\s+", " ", text)
    for abbr in ["e.g.", "i.e.", "et al.", "Fig.", "Eq.", "vs.", "ca.", "ref.", "Ref.", "al."]:
        text = text.replace(abbr, abbr.replace(".", "@@"))
    sents = re.split(r"(?<=[.!?])\s+(?=[A-Z])", text)
    return [s.replace("@@", ".").strip() for s in sents if len(s.split()) > 4]

def is_passive(s):
    return bool(re.search(r"\b(was|were|is|are|been|be|being)\s+\w+ed\b", s, re.IGNORECASE))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(
            int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        )
    return h

def add_table(doc, headers, rows, col_widths=None, header_color="1F3864"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        set_cell_bg(c, header_color)
        run = c.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        r = table.add_row()
        for i, val in enumerate(row):
            c = r.cells[i]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ── load & parse corpus ───────────────────────────────────────────────────────

print("Loading corpus...")
with open(CORPUS_PATH, encoding="utf-8") as f:
    raw = f.read()

papers = re.split(r"={40,}\n\[SOURCE:.*?\]\n={40,}", raw)
papers = [p.strip() for p in papers if len(p.strip()) > 200]
corpus_lower = raw.lower()

all_sents = []
passive_sents = []
active_sents = []

for paper in papers:
    sents = split_sentences(paper)
    for s in sents:
        all_sents.append(s)
        if is_passive(s):
            passive_sents.append(s)
        else:
            active_sents.append(s)

total = len(all_sents)
avg_len = sum(len(s.split()) for s in all_sents) / total

# ── Intro / Conclusion patterns ───────────────────────────────────────────────

INTRO_PHRASES = [
    "however", "to address", "this study", "in this study", "we report",
    "herein", "among the", "this work", "in this work", "despite", "although",
    "to overcome", "recently", "unfortunately", "have attracted", "nevertheless",
    "we demonstrate", "we propose", "is one of", "we present", "we develop",
    "considerable attention", "significant attention", "great attention",
    "over the past", "in recent years", "it has been",
]

CONCL_PHRASES = [
    "this work", "in summary", "overall", "in conclusion", "these results",
    "furthermore", "moreover", "in addition", "importantly", "notably",
    "we have", "we successfully", "we demonstrated", "we developed",
    "this provides", "this offers", "paving the way",
]

intro_counts = collections.Counter()
concl_counts = collections.Counter()

for paper in papers:
    lines = paper.split("\n")
    n = len(lines)
    intro_zone = " ".join(lines[:max(1, n // 4)]).lower()
    concl_zone = " ".join(lines[max(0, int(n * 0.80)):]).lower()
    for p in INTRO_PHRASES:
        c = intro_zone.count(p)
        if c:
            intro_counts[p] += c
    for p in CONCL_PHRASES:
        c = concl_zone.count(p)
        if c:
            concl_counts[p] += c

# ── Verb frequency ────────────────────────────────────────────────────────────

VERBS = [
    "shows", "showed", "exhibited", "exhibits", "improved", "enhanced",
    "enhances", "attributed", "increased", "reduced", "improve", "increase",
    "show", "enhance", "demonstrate", "demonstrated", "leads to", "achieved",
    "decreases", "decreased", "decrease", "reduces", "maintained", "delivered",
    "results in", "retain", "retained", "suppress", "suppressed", "achieve",
    "degrade", "degrades", "degraded", "deteriorate", "deteriorated",
    "fail", "fails", "failed", "exhibit", "demonstrates",
]

verb_counts = {}
for v in VERBS:
    c = len(re.findall(r"\b" + re.escape(v) + r"\b", corpus_lower))
    if c:
        verb_counts[v] = c

# ── Noun / phrase frequency ───────────────────────────────────────────────────

NOUNS = [
    "anode", "electrolyte", "separator", "sei", "cathode", "current collector",
    "binder", "conductivity", "lithiation", "sei layer", "si anode",
    "energy density", "eis", "active material", "capacity retention",
    "coulombic efficiency", "rate capability", "volume change", "areal capacity",
    "cycling stability", "delithiation", "impedance", "intercalation",
    "cycle stability", "ionic conductivity", "silicon anode", "volume expansion",
    "rate performance", "solid electrolyte interphase", "specific capacity",
    "charge transfer resistance", "pulverization", "initial coulombic efficiency",
    "power density", "gravimetric capacity", "agglomeration",
]

noun_counts = {}
for n in NOUNS:
    c = corpus_lower.count(n.lower())
    if c:
        noun_counts[n] = c

# ── Verb-noun collocations ────────────────────────────────────────────────────

KEY_VERBS = ["exhibit", "demonstrate", "show", "achieve", "deliver",
             "maintain", "retain", "improve", "enhance", "suppress",
             "attribute", "degrade", "deteriorate"]
KEY_NOUNS = ["capacity", "efficiency", "stability", "retention", "expansion",
             "resistance", "conductivity", "performance", "impedance", "sei"]

collocations = collections.Counter()
for s in all_sents:
    sl = s.lower()
    for v in KEY_VERBS:
        if re.search(r"\b" + v, sl):
            for n in KEY_NOUNS:
                if n in sl:
                    collocations[(v, n)] += 1

# ── Grammar pattern examples ──────────────────────────────────────────────────

passive_samples = [s for s in passive_sents if 12 < len(s.split()) < 30][:10]
active_samples  = [s for s in active_sents  if 12 < len(s.split()) < 28][:10]

# ── Sentence starters ─────────────────────────────────────────────────────────

STARTERS = {
    "The + noun + verb": r"^The \w+ \w+(s|ed|ing)\b",
    "This + noun + verb": r"^This \w+ \w+(s|ed|ing)\b",
    "These + noun + verb": r"^These \w+ \w+(s|ed|ing)\b",
    "subject + is/are attributed to": r"\b(is|are) attributed to\b",
    "subject + exhibits/shows + noun": r"\b(exhibits|shows) (a|an|the)?\s*\w+",
    "Furthermore / Moreover": r"^(Furthermore|Moreover),",
    "However,": r"^However,",
    "In this work/study,": r"^In this (work|study),",
    "Herein,": r"^Herein,",
    "To address/overcome": r"^To (address|overcome|tackle)",
}

starter_counts = {}
for label, pattern in STARTERS.items():
    c = sum(1 for s in all_sents if re.search(pattern, s, re.IGNORECASE))
    starter_counts[label] = c

# ── Degradation word comparison ───────────────────────────────────────────────

DEGRADE_WORDS = [
    ("degrade / degraded / degradation", r"\bdegrad\w*\b"),
    ("deteriorate / deteriorated",        r"\bdeteriorat\w*\b"),
    ("decay / decayed",                   r"\bdecay\w*\b"),
    ("fade / faded (as verb)",            r"\bfade[sd]?\b"),
    ("fail / failed / failure",           r"\bfail\w*\b"),
]
degrade_counts = []
for label, pattern in DEGRADE_WORDS:
    c = len(re.findall(pattern, corpus_lower))
    degrade_counts.append((label, c))

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

print("Building Word document...")
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("논문 영어 작성 스타일 분석 보고서", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.size = Pt(18)

sub = doc.add_paragraph("Si/C 복합 음극재 · 리튬이온배터리 논문 49편 (13,159 문장) 분석")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(11)
sub.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — 문장 구조
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. 문장 구조 통계", 1)

add_table(doc,
    headers=["항목", "수치"],
    rows=[
        ["분석 문장 수",        f"{total:,} 문장"],
        ["평균 문장 길이",      f"{avg_len:.1f} 단어/문장"],
        ["수동태 문장 수",      f"{len(passive_sents):,} 문장"],
        ["능동태 문장 수",      f"{len(active_sents):,} 문장"],
        ["수동태 비율",         f"{100*len(passive_sents)/total:.1f} %"],
        ["능동태 비율",         f"{100*len(active_sents)/total:.1f} %"],
    ],
    col_widths=[8, 6],
)
doc.add_paragraph()

# ── 수동태 예문
add_heading(doc, "1-1. 수동태 예문 (피동 구문 샘플)", 2)
doc.add_paragraph("수동태는 실험 절차 기술, 원인 귀속(attributed to), 선행 연구 인용에 집중적으로 사용됩니다.")
for s in passive_samples:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(s)
    run.font.size = Pt(10)
    run.font.italic = True

# ── 능동태 예문
add_heading(doc, "1-2. 능동태 예문 (결과 보고 샘플)", 2)
doc.add_paragraph("결과 보고, 기여점 선언, 실험 관찰에는 능동태가 압도적으로 사용됩니다.")
for s in active_samples:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(s)
    run.font.size = Pt(10)
    run.font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — 문장 시작 패턴
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. 빈출 문장 시작 패턴 (Grammar Structures)", 1)

add_table(doc,
    headers=["패턴", "출현 횟수", "설명"],
    rows=[
        [label, f"{cnt:,}", ""]
        for label, cnt in sorted(starter_counts.items(), key=lambda x: -x[1])
    ],
    col_widths=[7, 4, 5.5],
)
doc.add_paragraph()

add_heading(doc, "핵심 문법 구조 패턴", 2)
patterns = [
    ("결과 보고 (능동)",   "The [electrode/material] exhibits/shows [a + 형용사 + 명사]."),
    ("결과 보고 (능동)",   "[Figure N] shows [the + 명사 + of + 명사]."),
    ("원인 귀속 (수동)",   "The improved [성능] is attributed to [the + 명사]."),
    ("인과 관계",          "[현상] results in / leads to [결과]."),
    ("서론 문제 제시",     "However, [문제점]. / Despite [명사구], [문제점]."),
    ("서론 해결책 선언",   "To address/overcome this [challenge/issue], [we + 동사]."),
    ("서론 기여점 선언",   "In this work/study, we report/demonstrate/propose [명사구]."),
    ("결론 도입",          "In this work/study, we have demonstrated/developed [명사구]."),
    ("결론 추가 강조",     "Furthermore/Moreover/Importantly, [결과 수치]."),
]
add_table(doc,
    headers=["용도", "패턴 구조"],
    rows=patterns,
    col_widths=[4.5, 12],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — 전환어 (Transition Words)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. 빈출 전환어 (Transition Words)", 1)

add_heading(doc, "3-1. 서론부 전환어 (Introduction)", 2)
add_table(doc,
    headers=["전환어 / 구문", "출현 횟수", "사용 위치"],
    rows=[
        [p, f"{c:,}", "서론"]
        for p, c in intro_counts.most_common(20)
    ],
    col_widths=[6, 3, 3],
)
doc.add_paragraph()

add_heading(doc, "3-2. 결론부 전환어 (Conclusion)", 2)
add_table(doc,
    headers=["전환어 / 구문", "출현 횟수", "사용 위치"],
    rows=[
        [p, f"{c:,}", "결론"]
        for p, c in concl_counts.most_common(15)
    ],
    col_widths=[6, 3, 3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — 빈출 동사
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. 빈출 동사 순위 (Verb Frequency)", 1)

add_heading(doc, "4-1. 결과 보고 동사 (상위 30위)", 2)
verb_rows = sorted(verb_counts.items(), key=lambda x: -x[1])[:30]
add_table(doc,
    headers=["순위", "동사", "출현 횟수"],
    rows=[[str(i+1), v, f"{c:,}"] for i, (v, c) in enumerate(verb_rows)],
    col_widths=[2, 7, 4],
)
doc.add_paragraph()

add_heading(doc, "4-2. 용도별 동사 선호 순위", 2)
verb_pref = [
    ("결과 보고",   "show > exhibit > demonstrate > deliver",  "show 672회, exhibit 498회, demonstrate 165회"),
    ("성능 향상",   "improve > enhance > increase",            "improve 493회, enhance 325회, increase 304회"),
    ("성능 감소",   "reduce > decrease > suppress",            "reduce 204회, decrease 120회"),
    ("열화 표현",   "degrade > fail > deteriorate > decay > fade", "degrade 109회 >> deteriorate 27회"),
    ("원인 귀속",   "is attributed to",                       "203회 (필수 표현)"),
    ("인과 관계",   "leads to / results in",                  "leads to 69회, results in 46회"),
    ("성능 달성",   "achieve / maintain / retain",             "전기화학 수치 기술에 사용"),
]
add_table(doc,
    headers=["용도", "선호 순서", "코퍼스 근거"],
    rows=verb_pref,
    col_widths=[3.5, 6.5, 6.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — 빈출 명사/전문 용어
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. 빈출 명사 / 전문 용어 (Noun Frequency)", 1)

add_heading(doc, "5-1. 전체 출현 빈도 순위", 2)
noun_rows = sorted(noun_counts.items(), key=lambda x: -x[1])
add_table(doc,
    headers=["순위", "명사 / 구문", "출현 횟수"],
    rows=[[str(i+1), n, f"{c:,}"] for i, (n, c) in enumerate(noun_rows)],
    col_widths=[2, 9, 4],
)
doc.add_paragraph()

add_heading(doc, "5-2. 카테고리별 필수 전문 용어", 2)
term_cats = [
    ("구성 요소",     "anode, cathode, electrolyte, separator, current collector, binder, active material, SEI / SEI layer"),
    ("전기화학 성능", "capacity retention, cycling stability, Coulombic efficiency, initial Coulombic efficiency (ICE), rate capability, specific capacity, areal capacity, energy density, power density"),
    ("열화 메커니즘", "volume expansion, volume change, pulverization, cracking, fracture, agglomeration, capacity fade (명사)"),
    ("전기화학 공정", "lithiation, delithiation, intercalation, deintercalation"),
    ("저항 분석",     "charge transfer resistance, impedance (EIS), ionic conductivity, electronic conductivity, diffusion coefficient"),
]
add_table(doc,
    headers=["카테고리", "필수 용어"],
    rows=term_cats,
    col_widths=[4, 12.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — 동사+명사 콜로케이션
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. 빈출 동사-명사 콜로케이션 (Verb-Noun Collocations)", 1)

doc.add_paragraph("논문에서 가장 자주 함께 사용되는 동사-명사 조합입니다. 이 표현들을 그대로 사용하면 자연스러운 학술 문체가 됩니다.")
doc.add_paragraph()

top_colloc = collocations.most_common(25)
add_table(doc,
    headers=["순위", "동사", "명사", "출현 횟수", "예시 구문"],
    rows=[
        [str(i+1), v, n, f"{c:,}", f"...{v}s a high {n}..."]
        for i, ((v, n), c) in enumerate(top_colloc)
    ],
    col_widths=[1.5, 3, 3.5, 3, 5.5],
)
doc.add_paragraph()

add_heading(doc, "권장 표현 (Copy-and-Use)", 2)
copy_phrases = [
    ("용량 보고",     "exhibits a specific capacity of X mAh g⁻¹ at Y C"),
    ("용량 보고",     "delivers a reversible capacity of X mAh g⁻¹"),
    ("성능 향상",     "improves the overall electrochemical performance"),
    ("안정성 향상",   "improves the cycling stability significantly"),
    ("전도성 향상",   "enhances the ionic/electronic conductivity"),
    ("부피 팽창 억제","suppresses the volume expansion of Si"),
    ("원인 귀속",     "is attributed to the improved SEI stability"),
    ("결론 도입",     "In this work, we have demonstrated/developed ..."),
    ("결론 추가 강점","Furthermore, the electrode exhibits ... after N cycles."),
    ("서론 전환",     "However, [문제점]. To address this challenge, ..."),
]
add_table(doc,
    headers=["용도", "권장 표현"],
    rows=copy_phrases,
    col_widths=[4, 12.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — 열화 표현 비교
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. 열화 표현 선호도 비교 (Critical Rule)", 1)

doc.add_paragraph(
    "교수님은 열화 표현으로 'degrade/degradation'을 압도적으로 선호합니다. "
    "AI 번역기나 일반 작성 시 자주 나오는 'deteriorate', 'decay', 'fade'는 "
    "반드시 아래 표에 따라 교정해야 합니다."
)
doc.add_paragraph()

add_table(doc,
    headers=["표현", "출현 횟수", "선호도", "교정 방향"],
    rows=[
        ["degrade / degraded / degradation", f"{degrade_counts[0][1]:,}", "★★★★★ 1순위", "→ 그대로 사용"],
        ["fail / failed / failure",          f"{degrade_counts[4][1]:,}", "★★★★  2순위", "→ 그대로 사용"],
        ["deteriorate / deteriorated",       f"{degrade_counts[1][1]:,}", "★★★   3순위", "→ degrade로 교체 권장"],
        ["decay / decayed",                  f"{degrade_counts[2][1]:,}", "★★    4순위", "→ degrade로 교체 권장"],
        ["fade / faded (동사 용법)",          f"{degrade_counts[3][1]:,}", "★     피해야 함", "→ degrade로 교체"],
    ],
    col_widths=[6, 3.5, 4, 4],
)

doc.add_paragraph()
doc.add_paragraph("※ 'capacity fade'는 명사 형태로만 허용 (동사 'faded' 사용 금지)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — 표기 규칙
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. 표기 및 형식 규칙 (Formatting Rules)", 1)

format_rules = [
    ("Figure 표기",    "본문: Figure 2a (전체 표기)  /  괄호 내: (Fig. 2a)"),
    ("단위 표기",      "값과 단위 사이 공백: 1200 mAh g⁻¹  /  0.5 A g⁻¹  /  1.0 C"),
    ("소수점 구분자",  "쉼표(,) 아닌 마침표(.) 사용: 89.4 %"),
    ("퍼센트 표기",    "공백 없이 붙여씀: 89%  (89 % 아님)"),
    ("싸이클 수 표기", "10 이하: 영어 표기 (five cycles)  /  10 초과: 숫자 (200 cycles)"),
    ("대문자 규칙",    "Coulombic efficiency  /  Figure  /  Table (항상 대문자)"),
    ("화학식 표기",    "Li⁺, TiO₂, SiOₓ, Li₄Ti₅O₁₂ (유니코드 또는 LaTeX)"),
    ("1인칭 표기",    "'we' 사용 (the authors / the present authors 금지)"),
    ("축약어 금지",    "it's, don't 등 축약어 사용 금지 (학술 문체)"),
]
add_table(doc,
    headers=["규칙", "상세 내용"],
    rows=format_rules,
    col_widths=[4, 12.5],
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph(f"분석 논문 수: 49편  |  분석 문장 수: {total:,}  |  출처: merged_corpus.txt")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(0x90, 0x90, 0x90)

doc.save(OUTPUT_PATH)
print(f"\n완료! 저장 위치: {OUTPUT_PATH}")
